from io import BytesIO

from docx import Document

from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer
)

from reportlab.lib.styles import (
    getSampleStyleSheet
)

from reportlab.lib.pagesizes import (
    letter
)


# =====================================
# TXT EXPORT
# =====================================

def export_to_txt(content):

    return content.encode("utf-8")


# =====================================
# DOCX EXPORT
# =====================================

def export_to_docx(content):

    doc = Document()

    doc.add_heading(
        "Exported Content",
        level=1
    )

    doc.add_paragraph(content)

    buffer = BytesIO()

    doc.save(buffer)

    buffer.seek(0)

    return buffer


# =====================================
# PDF EXPORT
# =====================================

def export_to_pdf(content):

    buffer = BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter
    )

    styles = getSampleStyleSheet()

    elements = []

    elements.append(
        Paragraph(
            content,
            styles["BodyText"]
        )
    )

    elements.append(
        Spacer(1, 12)
    )

    doc.build(elements)

    buffer.seek(0)

    return buffer